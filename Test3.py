from openai import OpenAI
import asyncio
import logging
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.methods import DeleteWebhook
from aiogram.types import Message
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
import requests
import base64
import PyPDF2
from docx import Document
import io

# Настройка OpenAI
api_key = "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"
model = "openrouter/bert-nebulon-alpha"

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=api_key,
)

TOKEN = '8206219495:AAGFmDTDUtPtalFgR1slkGvE2hqMeCnKQN8'

logging.basicConfig(level=logging.INFO)
bot = Bot(TOKEN)
dp = Dispatcher()

@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    await message.answer(
        'Привет! Я короч мега крутой генератор с подключенной нейросетью привет. Используй команду /analyze_file для анализа файлов',
        parse_mode='HTML')

class FileStates(StatesGroup):
        waiting_for_file = State()

async def send_typing_animation(chat_id: int):
    """Показываем анимацию печати"""
    await bot.send_chat_action(chat_id, "typing")

async def send_animated_message(chat_id: int, text: str, delay: float = 0.5):
    """Отправляем сообщение с анимацией появления текста"""
    message = await bot.send_message(chat_id, "⏳ " + text)
    return message

async def edit_message_with_animation(message: types.Message, new_text: str, delay: float = 0.1):
    """Плавно редактируем сообщение с анимацией"""
    try:
            # Если текст слишком длинный, разбиваем на части
        if len(new_text) > 4000:
            parts = []
            current_part = ""

            for line in new_text.split('\n'):
                if len(current_part + line) < 4000:
                    current_part += line + '\n'
                else:
                    parts.append(current_part)
                    current_part = line + '\n'

            if current_part:
                    parts.append(current_part)

                # Редактируем первое сообщение
            await message.edit_text(parts[0])

                # Отправляем остальные части как новые сообщения
            for part in parts[1:]:
                await message.answer(part)

            return

            # Для коротких текстов - плавное редактирование
        current_text = ""
        words = new_text.split()

        for i, word in enumerate(words):
            current_text += word + " "

                # Обновляем каждые 5 слов или если это последнее слово
            if i % 5 == 0 or i == len(words) - 1:
                try:
                    await message.edit_text(current_text + "▌")
                    await asyncio.sleep(delay)
                except:
                    continue

            # Финальное обновление без курсора
        await message.edit_text(new_text)

    except Exception as e:
            # Если редактирование не удалось, отправляем новое сообщение
        await message.answer(new_text)

class FileStates(StatesGroup):
    waiting_for_file = State()

@dp.message(Command("analyze_file"))
async def start_file_analysis(message: Message, state: FSMContext):
    await message.answer(
        "Пожалуйста, отправьте файл для анализа. Поддерживаются:\n"
        "📷 Фото (JPEG, PNG)\n"
        "📄 Документы (PDF, TXT, DOCX)"
    )
    await state.set_state(FileStates.waiting_for_file)

async def analyze_file(file_data: bytes, file_type: str, file_name: str) -> str:

    if file_type == 'pdf':
        return await analyze_pdf(file_data)
    elif file_type == 'text':
        return await analyze_text(file_data)
    elif file_type == 'docx':
        return await analyze_docx(file_data)
    elif file_type == 'image':
        return await analyze_image(file_data)
    else:
        return "❌ Этот тип файла пока не поддерживается"
@dp.message(FileStates.waiting_for_file, F.document | F.photo)
async def handle_file_analysis(message: Message, state: FSMContext):
    try:
        file_data = None
        file_name = ""
        file_type = ""

        # Определяем тип файла и получаем данные
        if message.document:
            document = message.document
            file_name = document.file_name or "document"
            file_type = get_file_type(file_name)
            file_data = await download_file(document.file_id)

        elif message.photo:
            photo = message.photo[-1]
            file_name = "photo.jpg"
            file_type = "image"
            file_data = await download_file(photo.file_id)

        if not file_data:
            await message.answer("❌ Не удалось загрузить файл")
            return
        await send_typing_animation(message.chat.id)

        # Отправляем сообщение "Генерация..."
        status_message = await send_animated_message(message.chat.id, "Генерация ответа...")

        # Анализируем в зависимости от типа файла
        analysis_result = await analyze_file(file_data, file_type, file_name)

        # Плавно заменяем "Генерация..." на результат
        await edit_message_with_animation(status_message, analysis_result)

        await state.clear()

    except Exception as e:
        await message.answer(f"❌ Ошибка: {str(e)}")
        await state.clear()

def get_file_type(file_name: str) -> str:
    """Определяем тип файла по расширению"""
    ext = file_name.lower().split('.')[-1]

    if ext in ['pdf']:
        return 'pdf'
    elif ext in ['txt', 'log', 'md']:
        return 'text'
    elif ext in ['docx', 'doc']:
        return 'docx'
    elif ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
        return 'image'
    else:
        return 'unknown'

async def download_file(file_id: str) -> bytes:
    """Скачиваем файл из Telegram"""
    file_info = await bot.get_file(file_id)
    file_path = file_info.file_path
    file_url = f"https://api.telegram.org/file/bot{TOKEN}/{file_path}"
    response = requests.get(file_url)
    return response.content

async def analyze_pdf(file_data: bytes) -> str:
    """Анализ PDF файла с помощью PyPDF2"""
    try:
        # Используем BytesIO для работы с данными в памяти
        pdf_file = io.BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Извлекаем текст
        full_text = ""
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    full_text += f"--- Страница {page_num + 1} ---\n{page_text}\n\n"
            except Exception as e:
                continue  # Пропускаем страницы с ошибками

        if full_text.strip():
            return await analyze_with_ai(full_text, "PDF документ")
        else:
            return "❌ В PDF не найден текст или документ состоит из сканированных изображений"

    except Exception as e:
        return f"❌ Ошибка анализа PDF: {str(e)}"

async def analyze_text(file_data: bytes) -> str:
    """Анализ текстового файла"""
    try:
        text_content = file_data.decode('utf-8', errors='ignore')

        if text_content.strip():
            return await analyze_with_ai(text_content, "текстовый файл")
        else:
            return "❌ Файл пуст"

    except Exception as e:
        return f"❌ Ошибка анализа текста: {str(e)}"

async def analyze_docx(file_data: bytes) -> str:
    """Анализ DOCX файла"""
    try:
        # Используем BytesIO для работы с данными в памяти
        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)
        full_text = ""

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + "\n"

        if full_text.strip():
            return await analyze_with_ai(full_text, "Word документ")
        else:
            return "❌ Документ пуст"

    except Exception as e:
        return f"❌ Ошибка анализа DOCX: {str(e)}"

async def analyze_image(file_data: bytes) -> str:
    """Анализ изображения"""
    try:
        image_base64 = base64.b64encode(file_data).decode('utf-8')

        chat_response = client.chat.completions.create(
            model="openrouter/bert-nebulon-alpha",
            messages=[
                {
                    "role": "system",
                    "content": "Кратко опиши что изображено на картинке",
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Что изображено на этой картинке?"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}"
                            }
                        }
                    ]
                },
            ],
            extra_headers={
                "HTTP-Referer": "https://your-site.com",
                "X-Title": "Telegram Bot"
            }
        )

        return f"🖼️ {chat_response.choices[0].message.content}"

    except Exception as e:
        return f"❌ Ошибка анализа изображения: {str(e)}"

async def analyze_with_ai(content: str, file_type: str) -> str:
    """Общий анализ контента с помощью нейросети"""
    try:
        # Ограничиваем длину
        if len(content) > 6000:
            content = content[:6000] + "..."

        chat_response = client.chat.completions.create(
            model="openrouter/bert-nebulon-alpha",
            messages=[
                {
                    "role": "system",
                    "content": "Кратко проанализируй содержание и выдели основное",
                },
                {
                    "role": "user",
                    "content": f"Кратко опиши содержание этого {file_type}:\n\n{content}"
                },
            ],
            extra_headers={
                "HTTP-Referer": "https://your-site.com",
                "X-Title": "Telegram Bot"
            }
        )

        return chat_response.choices[0].message.content
    except Exception as e:
        return f"❌ Ошибка анализа ИИ: {str(e)}"

@dp.message(FileStates.waiting_for_file, Command("cancel"))
async def cancel_file_analysis(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Анализ файла отменен")

async def main():
    await bot(DeleteWebhook(drop_pending_updates=True))
    await dp.start_polling(bot)

if __name__ == "__main__":

    asyncio.run(main())
